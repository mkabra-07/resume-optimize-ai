"""
DOCX Exporter — Generates a clean, ATS-friendly DOCX from structured ResumeContent.

Sections: Contact, Summary, Skills, Experience, Projects, Education, Certifications.
Formatting rules:
  - No tables, images, or multi-column layouts (ATS-friendly).
  - Clear heading hierarchy.
  - Bullet points for experience and project entries.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from schemas.resume import ResumeContent


def _add_horizontal_rule(doc: Document) -> None:
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_heading_style(para, size: int = 13, bold: bool = True, color: str = "1F3864") -> None:
    """Apply consistent heading formatting."""
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)


def export_docx(content: ResumeContent, version_name: str, output_path: str) -> str:
    """
    Generate an ATS-friendly DOCX resume from structured content.

    Args:
        content:     The structured ResumeContent to render.
        version_name: Used for internal reference only (not written to doc).
        output_path: Absolute path where the .docx file will be saved.

    Returns:
        The output_path on success.
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()

    # --- Page Margins (standard 1-inch) ---
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # ── CONTACT ───────────────────────────────────────────────────────────────
    contact = content.contact
    if contact.name:
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(contact.name)
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_parts = []
    if contact.email:
        contact_parts.append(contact.email)
    if contact.phone:
        contact_parts.append(contact.phone)
    if contact.location:
        contact_parts.append(contact.location)
    if contact.linkedin:
        contact_parts.append(contact.linkedin)
    if contact.github:
        contact_parts.append(contact.github)

    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(10)

    _add_horizontal_rule(doc)

    # ── SUMMARY ───────────────────────────────────────────────────────────────
    if content.summary:
        h = doc.add_heading("PROFESSIONAL SUMMARY", level=1)
        _set_heading_style(h)
        doc.add_paragraph(content.summary)
        _add_horizontal_rule(doc)

    # ── SKILLS ────────────────────────────────────────────────────────────────
    if content.skills:
        h = doc.add_heading("SKILLS", level=1)
        _set_heading_style(h)
        # Group into lines of ~6 skills to stay ATS-readable without tables
        chunk_size = 8
        chunks = [content.skills[i : i + chunk_size] for i in range(0, len(content.skills), chunk_size)]
        for chunk in chunks:
            doc.add_paragraph("  •  ".join(chunk))
        _add_horizontal_rule(doc)

    # ── EXPERIENCE ────────────────────────────────────────────────────────────
    if content.experience:
        h = doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
        _set_heading_style(h)
        for exp in content.experience:
            # Job title + company line
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{exp.title}")
            title_run.font.bold = True
            title_run.font.size = Pt(11)
            title_para.add_run(f"  —  {exp.company}")

            # Date + location line
            date_str = f"{exp.start_date} – {exp.end_date}"
            if exp.location:
                date_str += f"  |  {exp.location}"
            date_para = doc.add_paragraph(date_str)
            for run in date_para.runs:
                run.font.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            for bullet in exp.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

            doc.add_paragraph()  # spacer

        _add_horizontal_rule(doc)

    # ── PROJECTS ──────────────────────────────────────────────────────────────
    if content.projects:
        h = doc.add_heading("PROJECTS", level=1)
        _set_heading_style(h)
        for proj in content.projects:
            proj_para = doc.add_paragraph()
            name_run = proj_para.add_run(proj.name)
            name_run.font.bold = True
            if proj.technologies:
                proj_para.add_run(f"  |  {', '.join(proj.technologies)}")
            if proj.url:
                proj_para.add_run(f"  ({proj.url})")
            if proj.description:
                doc.add_paragraph(proj.description)
            for bullet in proj.bullets:
                doc.add_paragraph(bullet, style="List Bullet")
            doc.add_paragraph()

        _add_horizontal_rule(doc)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    if content.education:
        h = doc.add_heading("EDUCATION", level=1)
        _set_heading_style(h)
        for edu in content.education:
            edu_para = doc.add_paragraph()
            deg_run = edu_para.add_run(f"{edu.degree}")
            deg_run.font.bold = True
            if edu.field:
                edu_para.add_run(f" in {edu.field}")

            inst_para = doc.add_paragraph(edu.institution)
            dates = []
            if edu.start_date:
                dates.append(edu.start_date)
            if edu.end_date:
                dates.append(edu.end_date)
            if dates:
                inst_para.add_run(f"  |  {' – '.join(dates)}")
            if edu.gpa:
                inst_para.add_run(f"  |  GPA: {edu.gpa}")
            doc.add_paragraph()

        _add_horizontal_rule(doc)

    # ── CERTIFICATIONS ────────────────────────────────────────────────────────
    if content.certifications:
        h = doc.add_heading("CERTIFICATIONS", level=1)
        _set_heading_style(h)
        for cert in content.certifications:
            cert_parts = [cert.name]
            if cert.issuer:
                cert_parts.append(cert.issuer)
            if cert.date:
                cert_parts.append(cert.date)
            doc.add_paragraph("  •  ".join(cert_parts))

    doc.save(output_path)
    return output_path
